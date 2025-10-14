import os
import json
import hashlib
from datetime import datetime
from typing import List, Dict, Tuple, Optional

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def _safe_filename(name: str) -> str:
    bad = '<>:"/\\|?*'
    safe = "".join("_" if ch in bad else ch for ch in name).strip()
    return safe or "Test"


def _image_file_info(path: Optional[str]) -> Tuple[str, int, float]:
    if not path or not isinstance(path, str):
        return ("", 0, 0.0)
    try:
        stat = os.stat(path)
        return (path, int(stat.st_size), float(stat.st_mtime))
    except Exception:
        return (path, 0, 0.0)


def _calc_questions_hash(questions: List[Dict]) -> str:
    norm = []
    for q in questions:
        q_text = str(q.get("question", "")).strip()
        answers = q.get("answers", []) or []
        answers_norm = [(str(a.get("text", "")).strip(), bool(a.get("correct", False))) for a in answers]
        img_info = _image_file_info(q.get("image"))
        norm.append({
            "q": q_text,
            "a": answers_norm,
            "img": img_info,
        })
    payload = json.dumps(norm, ensure_ascii=False, sort_keys=True, separators=(",", ":"))
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _write_meta(meta_path: str, meta: Dict):
    try:
        with open(meta_path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def _read_meta(meta_path: str) -> Dict:
    if not os.path.exists(meta_path):
        return {}
    try:
        with open(meta_path, "r", encoding="utf-8") as f:
            data = json.load(f)
            return data if isinstance(data, dict) else {}
    except Exception:
        return {}


def _build_docx(test_name: str, questions: List[Dict], output_path: str):
    """
    Формує DOCX у форматі твоїх скриптів:
    - Заголовок = назва тесту
    - Питання абзацом
    - Картинка (якщо є) 4", по центру
    - Відповіді кожна абзацом; правильна — жирним
    - Порожній рядок після питання
    """
    doc = Document()
    doc.add_heading(test_name, level=0)

    for idx, q in enumerate(questions, start=1):
        q_text = str(q.get("question", "")).strip()
        answers = q.get("answers", []) or []

        # Питання
        doc.add_paragraph(q_text, style="Normal")

        # Картинка з поля 'image' (як у бота після attach_images)
        img_path = q.get("image")
        if isinstance(img_path, str) and os.path.exists(img_path):
            try:
                doc.add_picture(img_path, width=Inches(4))
                last_par = doc.paragraphs[-1]
                last_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass

        # Відповіді
        for ans in answers:
            text = str(ans.get("text", "")).strip()
            correct = bool(ans.get("correct", False))
            p = doc.add_paragraph()
            run = p.add_run(text)
            if correct:
                run.bold = True

        # Відступ
        doc.add_paragraph("")

    doc.save(output_path)


def export_test_to_docx(test_name: str, questions: List[Dict], output_dir: str = "tests") -> Tuple[str, bool]:
    """
    Експортує тест у DOCX.
    Повертає (path_to_docx, regenerated_flag).

    Логіка:
    - Рахуємо хеш вмісту (питання/відповіді/картинки).
    - Якщо існуючий DOCX + .meta.json з таким самим хешем — повертаємо існуючий (regenerated=False).
    - Інакше — генеруємо поверх (regenerated=True) і оновлюємо .meta.json.
    """
    os.makedirs(output_dir, exist_ok=True)

    safe_name = _safe_filename(test_name)
    docx_path = os.path.join(output_dir, f"{safe_name}.docx")
    meta_path = os.path.join(output_dir, f"{safe_name}.docx.meta.json")

    content_hash = _calc_questions_hash(questions)
    meta = _read_meta(meta_path)
    prev_hash = meta.get("content_hash")
    doc_exists = os.path.exists(docx_path)

    if doc_exists and prev_hash == content_hash:
        return docx_path, False

    _build_docx(test_name, questions, docx_path)

    new_meta = {
        "test_name": test_name,
        "content_hash": content_hash,
        "question_count": len(questions),
        "generated_at": datetime.utcnow().isoformat() + "Z",
        "docx_path": docx_path,
    }
    _write_meta(meta_path, new_meta)

    return docx_path, True
